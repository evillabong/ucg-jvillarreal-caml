from pathlib import Path
import shutil

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parents[1]
SOURCE_DIR = BASE_DIR / "informe ejecutivo"
TARGET_DIR = BASE_DIR / "informe ejecutivo v2"
IMAGES_DIR = TARGET_DIR / "imagenes"
OUTPUT_DOCX = TARGET_DIR / "Informe_Ejecutivo_Detallado_Desercion_Escolar_v2.docx"

IMAGES = [
    "dist_target.png",
    "desercion_enums.png",
    "proyeccion_desercion_regiones.png",
    "inicio_fin_region.png",
    "brechas_factores_regionales.png",
    "heatmap_provincia_anio.png",
    "feature_importance.png",
    "correlacion.png",
    "confusion_matrix.png",
    "pca_desercion.png",
]

METRICS = {
    "periodo": "2015-2016 a 2024-2025",
    "n_anios": 10,
    "registros": 163456,
    "features": 16,
    "train": 130764,
    "test": 32692,
    "accuracy": 0.6506790652147314,
    "f1_weighted": 0.7235107338715084,
    "cv_mean": 0.7173049119118452,
    "cv_std": 0.0027173867697181183,
    "bajo": 147118,
    "medio": 10200,
    "alto": 6138,
}

TOP_FEATURES = [
    ("est_inicio", 0.183793),
    ("ratio_doc_est", 0.169449),
    ("ratio_genero_est", 0.168826),
    ("total_docentes", 0.114938),
    ("zona", 0.083301),
]

TOP_PROVINCES = [
    ("Napo", 4.68),
    ("Pichincha", 3.44),
    ("Pastaza", 3.38),
    ("Sucumbios", 3.34),
    ("Azuay", 2.99),
]

REGIONAL_SUMMARY = [
    ("Sierra", 1.90, 1.12, -0.36, 0.94, -0.51),
    ("Costa", 1.48, 1.19, -0.44, 0.62, -0.49),
]


def set_cell_text(cell, text, bold=False):
    cell.text = str(text)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            run.bold = bold


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def format_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Calibri"
    styles["Title"].font.size = Pt(24)
    styles["Title"].font.bold = True
    styles["Heading 1"].font.name = "Calibri"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = RGBColor(31, 78, 121)
    styles["Heading 2"].font.name = "Calibri"
    styles["Heading 2"].font.size = Pt(12.5)
    styles["Heading 2"].font.bold = True
    styles["Heading 2"].font.color.rgb = RGBColor(55, 96, 146)


def add_title_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("INFORME EJECUTIVO DETALLADO")
    r.bold = True
    r.font.size = Pt(24)
    r.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Prediccion de Desercion Escolar en Ecuador")
    r.bold = True
    r.font.size = Pt(20)
    r.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(
        "Analisis de Machine Learning sobre datos AMIE del Ministerio de Educacion del Ecuador"
    ).font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(
        "Periodo analizado: 2015-2016 a 2024-2025 | Elaborado en abril de 2026"
    ).italic = True

    doc.add_paragraph("")
    summary = doc.add_table(rows=4, cols=2)
    summary.alignment = WD_TABLE_ALIGNMENT.CENTER
    summary.style = "Table Grid"
    rows = [
        ("Registros analizados", f"{METRICS['registros']:,}"),
        ("Anios lectivos", str(METRICS["n_anios"])),
        ("Features del modelo", str(METRICS["features"])),
        ("Desempeno F1 weighted", f"{METRICS['f1_weighted']:.4f}"),
    ]
    for idx, (left, right) in enumerate(rows):
        set_cell_text(summary.rows[idx].cells[0], left, bold=True)
        set_cell_text(summary.rows[idx].cells[1], right)
        shade_cell(summary.rows[idx].cells[0], "D9EAF7")

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(
        "Este documento consolida los resultados del notebook desercion_escolar_ML.ipynb, "
        "integrando todas las visualizaciones generadas y una lectura ejecutiva de sus hallazgos. "
        "El objetivo es apoyar decisiones de politica educativa, focalizacion territorial y priorizacion "
        "de intervenciones de permanencia escolar."
    )

    doc.add_page_break()


def add_heading_paragraph(doc, heading, text):
    doc.add_heading(heading, level=1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(text)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, headers):
        set_cell_text(cell, text, bold=True)
        shade_cell(cell, "BDD7EE")
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            set_cell_text(cell, text)
    doc.add_paragraph("")


def add_figure(doc, image_name, caption, width=6.6):
    image_path = IMAGES_DIR / image_name
    if not image_path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(9.5)


def copy_images():
    IMAGES_DIR.mkdir(parents=True, exist_ok=True)
    for image_name in IMAGES:
        source = SOURCE_DIR / image_name
        target = IMAGES_DIR / image_name
        if source.exists():
            shutil.copy2(source, target)


def build_report():
    copy_images()

    doc = Document()
    format_document(doc)
    add_title_page(doc)

    add_heading_paragraph(
        doc,
        "1. Objetivo y alcance del ejercicio",
        "El ejercicio busca estimar y clasificar el nivel de desercion escolar de las instituciones educativas del Ecuador "
        "a partir de caracteristicas institucionales, geograficas y de recursos. Se procesaron datos oficiales AMIE de "
        "10 anios lectivos, unificando archivos de inicio y fin de anio por clave AMIE para calcular la tasa de desercion "
        "y generar un conjunto analitico consistente.",
    )
    add_table(
        doc,
        ["Indicador", "Valor"],
        [
            ("Periodo", METRICS["periodo"]),
            ("Anios lectivos", METRICS["n_anios"]),
            ("Registros institucion-anio", f"{METRICS['registros']:,}"),
            ("Features usadas", METRICS["features"]),
            ("Entrenamiento", f"{METRICS['train']:,}"),
            ("Prueba", f"{METRICS['test']:,}"),
        ],
    )

    add_heading_paragraph(
        doc,
        "2. Metodo de construccion del dataset y del modelo",
        "La canalizacion de datos normaliza encabezados heterogeneos entre periodos, tipifica variables mediante enums y "
        "construye una dataclass por institucion. Sobre el dataset final se entreno un Random Forest Classifier balanceado, "
        "con 200 arboles, profundidad maxima de 15 y validacion cruzada estratificada de 5 folds. La variable objetivo se "
        "clasifico en tres niveles: Bajo, Medio y Alto.",
    )
    add_table(
        doc,
        ["Metrica", "Resultado"],
        [
            ("Accuracy", f"{METRICS['accuracy']:.4f}"),
            ("F1 weighted", f"{METRICS['f1_weighted']:.4f}"),
            ("CV F1 promedio", f"{METRICS['cv_mean']:.4f}"),
            ("CV F1 desviacion", f"{METRICS['cv_std']:.4f}"),
        ],
    )

    add_heading_paragraph(
        doc,
        "3. Distribucion general del riesgo y evolucion temporal",
        "La muestra presenta una fuerte concentracion en instituciones con nivel Bajo de desercion. Esto impone un reto "
        "de desbalance de clases y obliga a interpretar el accuracy junto con el F1 weighted. Aun asi, el comportamiento "
        "historico permite distinguir cambios temporales claros y sostener el valor analitico del ejercicio.",
    )
    add_table(
        doc,
        ["Nivel", "Instituciones", "Participacion"],
        [
            ("Bajo", f"{METRICS['bajo']:,}", "90.00%"),
            ("Medio", f"{METRICS['medio']:,}", "6.24%"),
            ("Alto", f"{METRICS['alto']:,}", "3.76%"),
        ],
    )
    add_figure(
        doc,
        "dist_target.png",
        "Figura 1. Distribucion del nivel de desercion y evolucion temporal de la tasa media.",
    )
    doc.add_paragraph(
        "Lectura ejecutiva: el 90% de los registros cae en riesgo Bajo, pero el volumen absoluto de casos Medio y Alto sigue "
        "siendo material para una politica de alerta temprana, porque representa 16,338 observaciones en niveles que requieren atencion."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    add_heading_paragraph(
        doc,
        "4. Desercion segun categorias institucionales",
        "El analisis por enums muestra variaciones relevantes por sostenimiento, area, modalidad y jornada. Este bloque sirve "
        "para identificar grupos institucionales donde las tasas medias se apartan del comportamiento agregado y donde un piloto "
        "de intervencion puede producir mayor retorno.",
    )
    add_figure(
        doc,
        "desercion_enums.png",
        "Figura 2. Tasa media de desercion por sostenimiento, area, modalidad y jornada.",
    )

    add_heading_paragraph(
        doc,
        "5. Comparacion regional y magnitud de la perdida de matricula",
        "Las visualizaciones regionales ayudan a separar dos preguntas distintas: donde la tasa relativa es mayor y donde el volumen "
        "de estudiantes afectados es mas alto. El ultimo periodo disponible muestra a la Sierra con mayor tasa media, pero la Costa "
        "mantiene un volumen absoluto superior de matricula y, por tanto, un impacto agregado relevante.",
    )
    add_table(
        doc,
        [
            "Region",
            "Tasa ultimo periodo (%)",
            "Perdida inicio-fin (%)",
            "Rural-Urbana (p.p.)",
            "No fiscal-Fiscal (p.p.)",
            "Bajo ratio-Alto ratio (p.p.)",
        ],
        [
            (
                region,
                f"{tasa_ultimo:.2f}",
                f"{perdida:.2f}",
                f"{rural_urbana:.2f}",
                f"{nofiscal_fiscal:.2f}",
                f"{ratio_gap:.2f}",
            )
            for region, tasa_ultimo, perdida, rural_urbana, nofiscal_fiscal, ratio_gap in REGIONAL_SUMMARY
        ],
    )
    add_figure(
        doc,
        "proyeccion_desercion_regiones.png",
        "Figura 3. Proyeccion historica de la tasa media de desercion por region escolar.",
    )
    add_figure(
        doc,
        "inicio_fin_region.png",
        "Figura 4. Matricula acumulada por region al inicio y al fin de cada anio lectivo.",
    )
    add_figure(
        doc,
        "brechas_factores_regionales.png",
        "Figura 5. Brechas regionales en factores asociados a la desercion.",
    )
    doc.add_paragraph(
        "Lectura ejecutiva: la brecha mas consistente y accionable es No fiscal vs Fiscal, con diferencias de 0.62 p.p. en Costa "
        "y 0.94 p.p. en Sierra. Las brechas Rural-Urbana y Bajo ratio-Alto ratio son pequenas y cambian de signo en el agregado, "
        "por lo que deben tratarse como hallazgos descriptivos y no como evidencia causal concluyente."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    add_heading_paragraph(
        doc,
        "6. Concentracion territorial del riesgo",
        "La lectura provincial agrega granularidad a la priorizacion. Las provincias con mayor tasa promedio observada en el periodo "
        "fueron Napo, Pichincha, Pastaza, Sucumbios y Azuay. Este ranking debe cruzarse con volumen de matricula y capacidad operativa "
        "antes de asignar recursos, porque una tasa alta sobre una base pequena no implica necesariamente el mayor impacto total."
    )
    add_table(
        doc,
        ["Provincia", "Tasa media estimada (%)"],
        [(name, f"{value:.2f}") for name, value in TOP_PROVINCES],
    )
    add_figure(
        doc,
        "heatmap_provincia_anio.png",
        "Figura 6. Mapa de calor de la tasa de desercion por provincia y anio lectivo.",
        width=6.8,
    )

    add_heading_paragraph(
        doc,
        "7. Factores que mas pesan en la prediccion",
        "El modelo Random Forest ubica el tamano inicial de la matricula, la relacion docente-estudiante, la composicion por genero, "
        "el total de docentes y la zona como las variables mas informativas. Esta senal respalda la necesidad de combinar indicadores de "
        "escala institucional, dotacion de recursos y contexto territorial en cualquier estrategia de retencion.",
    )
    add_table(
        doc,
        ["Feature", "Importancia relativa"],
        [(name, f"{value:.6f}") for name, value in TOP_FEATURES],
    )
    add_figure(
        doc,
        "feature_importance.png",
        "Figura 7. Importancia relativa de variables en el modelo Random Forest.",
    )
    add_figure(
        doc,
        "correlacion.png",
        "Figura 8. Matriz de correlacion de variables numericas del modelo.",
        width=5.8,
    )

    add_heading_paragraph(
        doc,
        "8. Confiabilidad del modelo",
        "El accuracy del modelo es moderado, pero el F1 weighted es mas representativo por el desbalance de clases. La estabilidad en la "
        "validacion cruzada sugiere que el comportamiento general del modelo es consistente y no depende de una sola particion del dataset. "
        "Aun asi, la herramienta debe usarse como apoyo a decision y no como sustituto de criterio tecnico o levantamiento territorial."
    )
    add_figure(
        doc,
        "confusion_matrix.png",
        "Figura 9. Matriz de confusion del clasificador Random Forest sobre el conjunto de prueba.",
        width=5.8,
    )
    add_figure(
        doc,
        "pca_desercion.png",
        "Figura 10. Proyeccion PCA de las instituciones en dos componentes principales.",
        width=6.2,
    )

    add_heading_paragraph(
        doc,
        "9. Implicaciones para decision",
        "A nivel operativo, el ejercicio sugiere tres lineas de accion inmediatas. Primero, priorizar monitoreo diferenciado sobre instituciones "
        "no fiscales en Sierra y Costa, donde la brecha descriptiva es mas estable. Segundo, combinar lectura provincial con volumen de matricula "
        "para definir focalizacion territorial. Tercero, usar el modelo como sistema de alerta temprana anual para identificar instituciones de "
        "riesgo Medio y Alto antes del cierre del periodo."
    )
    add_table(
        doc,
        ["Horizonte", "Accion recomendada", "Justificacion"],
        [
            (
                "0-3 meses",
                "Implementar tablero de alerta temprana para riesgo Medio y Alto.",
                "Existen 16,338 observaciones historicas en niveles que requieren seguimiento.",
            ),
            (
                "3-6 meses",
                "Priorizar revision de instituciones no fiscales en Sierra y Costa.",
                "Es la brecha regional mas consistente del analisis descriptivo.",
            ),
            (
                "6-12 meses",
                "Integrar variables socioeconomicas y de conectividad.",
                "Permitiria mejorar explicacion causal y sensibilidad territorial del modelo.",
            ),
        ],
    )

    add_heading_paragraph(
        doc,
        "10. Limitaciones y cierre",
        "El analisis se basa en datos administrativos institucionales y no incorpora variables socioeconomicas del hogar, movilidad estudiantil ni "
        "factores locales de oferta educativa. Por ello, los resultados deben leerse como un instrumento de priorizacion analitica. Aun con esa "
        "limitacion, el ejercicio demuestra que los datos AMIE permiten construir una senal util para orientar decisiones de permanencia escolar."
    )

    footer_section = doc.sections[-1]
    footer = footer_section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = (
        "Informe generado desde desercion_escolar_ML.ipynb | Universidad de Guayaquil | CAML | Abril 2026"
    )

    doc.save(OUTPUT_DOCX)


if __name__ == "__main__":
    build_report()
    print(f"Informe generado en: {OUTPUT_DOCX}")